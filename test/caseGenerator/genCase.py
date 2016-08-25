#encoding:utf-8

import yaml
import sys
from docx import Document

class CaseGenerator(object):
    def __init__(self, inFile, outFile):
        self.document = Document(inFile)
        self.outfile = outFile
        self.cases = []

    def generate(self):
        self.__genCaseFromDocTable()
        self.__genYmlFile()
        
    def __genCaseFromDocTable(self):
        import types
        for table in self.document.tables:    
            case = {}
            id = self.__getCellText(table, 1, 1)
            case["name"] = self.__getCellText(table, 0, 1) + "_" + id    
            case["steps"] = []
    
            for row in range(3, len(table.rows)):
                step = {}
                step["connection"] = self.__getCellText(table, row, 1)
                if self.__getCellText(table, row, 2) == '<enter>':
                    step["command"] = ''
                else:
                    step["command"] = self.__getCellText(table, row, 2)
                result = yaml.load(self.__getCellText(table, row, 3)) 
                if type(result) == types.DictType:
                    for key in result:
                        if type(result[key]) == types.ListType:
                            for idx in range(0, len(result[key])):
                                result[key][idx] = str(result[key][idx])
                step["result"] = result
                case["steps"].append(step)           
            self.cases.append(case)   

    def __getCellText(self, table, row, cloumn):
        return eval(repr(table.cell(row,cloumn).paragraphs[0].text)[1:])

    def __genYmlFile(self):
        f = file(self.outfile, 'w')
        f.write(yaml.dump(self.cases))
        f.close()
    
if __name__=="__main__":
    print sys.argv[0], sys.argv[1]
    generator = CaseGenerator(sys.argv[1], sys.argv[2])   
    generator.generate()

